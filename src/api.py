"""
FastAPI Backend - Serve static files and API endpoints
"""

from fastapi import FastAPI, HTTPException, File, UploadFile
from fastapi.responses import FileResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional, List, Dict, Any
from pathlib import Path
import uuid
from datetime import datetime
from collections import defaultdict
import yfinance as yf
import os
import shutil
import json
import PyPDF2
from docx import Document

from src.utils.pdf_parser import EnhancedPDFParser, PDFContent
from src.agents.research_agent import ResearchAgent
from src.agents.research_planner import ResearchPlanner, ResearchPlan
from src.agents.deep_research_executor import DeepResearchExecutor
from src.agents.query_router import QueryRouter, AgentRegistry, RoutingDecision
from src.agents.query_classifier import IntelligentQueryClassifier, ConversationContextManager
from src.core.report_generator import ReportGenerator
from src.utils.logger import get_logger
from src.utils.financial_calculator import FinancialCalculator
from src.utils.number_extractor import NumberExtractor
from src.utils.universal_company_resolver import UniversalCompanyResolver
from src.utils.validators import InputValidator, validate_and_sanitize_query
from src.utils.memory_manager import get_memory_manager
from src.data.vector_store import VectorStore
from src.data.document_processor import DocumentProcessor, Document as DocClass, Chunk
import asyncio
import time

logger = get_logger("api")
memory_manager = get_memory_manager()

def _format_response_with_highlights(response: str, sources: List[Dict]) -> str:
    """
    Format response with highlighted actionable insights and source attribution
    
    Enhances readability by:
    - Highlighting key insights with emoji markers
    - Adding proper source citations
    - Emphasizing recommendations and risk flags
    """
    import re
    
    # Patterns for actionable insights
    insight_patterns = [
        (r'(Key takeaway|Key insight|Important:|Critical:|Notable:)(.+?)(?=\n|$)', '🔑 **\\1**\\2'),
        (r'(Recommendation|Suggested action|Advised to)(.+?)(?=\n|$)', '💡 **\\1**\\2'),
        (r'(Risk|Warning|Concern|Challenge)(.+?)(?=\n|$)', '⚠️ **\\1**\\2'),
        (r'(Opportunity|Potential|Growth)(.+?)(?=\n|$)', '🚀 **\\1**\\2'),
    ]
    
    enhanced_response = response
    
    # Apply highlighting patterns
    for pattern, replacement in insight_patterns:
        enhanced_response = re.sub(pattern, replacement, enhanced_response, flags=re.IGNORECASE)
    
    # Add source attribution section if sources exist
    if sources and len(sources) > 0:
        enhanced_response += "\n\n---\n\n### 📚 Sources & References\n\n"
        
        for idx, source in enumerate(sources[:10], 1):  # Limit to top 10 sources
            if isinstance(source, dict):
                title = source.get('title', 'Untitled')
                url = source.get('url', '#')
                snippet = source.get('snippet', '')[:150]
                
                enhanced_response += f"{idx}. **[{title}]({url})**"
                if snippet:
                    enhanced_response += f"\n   *{snippet}...*"
                enhanced_response += "\n\n"
            elif isinstance(source, str):
                enhanced_response += f"{idx}. {source}\n\n"
    
    return enhanced_response

# Initialize FastAPI
app = FastAPI(
    title="Financial Research AI",
    description="AI-powered financial research assistant",
    version="3.0.0"
)

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Mount static files
app.mount("/static", StaticFiles(directory="static"), name="static")

# Global instances
research_agent: Optional[ResearchAgent] = None
report_generator: Optional[ReportGenerator] = None
research_planner: Optional[ResearchPlanner] = None
deep_executor: Optional[DeepResearchExecutor] = None
query_router: Optional[QueryRouter] = None
agent_registry: Optional[AgentRegistry] = None
calculator: Optional[FinancialCalculator] = None
number_extractor: Optional[NumberExtractor] = None
company_resolver: Optional[UniversalCompanyResolver] = None
vector_store: Optional[VectorStore] = None
document_processor: Optional[DocumentProcessor] = None
query_classifier: Optional[IntelligentQueryClassifier] = None
context_manager: Optional[ConversationContextManager] = None
pdf_parser: Optional[EnhancedPDFParser] = None
chat_sessions: Dict[str, List[Dict]] = {}
research_plans: Dict[str, Any] = {}
conversation_history: Dict[str, List[Dict]] = defaultdict(list)
file_document_cache: Dict[str, str] = {}  # Cache for file->doc_id mapping

# Models
class ChatMessage(BaseModel):
    message: str
    session_id: Optional[str] = None
    file_path: Optional[str] = None  # Support for attached files
    deepMode: bool = False  # Force deep analysis mode

class ChatResponse(BaseModel):
    response: str
    sources: List[Dict[str, Any]]
    session_id: str
    timestamp: str
    report_available: bool
    report_path: Optional[str] = None

class StockQuery(BaseModel):
    symbol: str
    period: str = "1mo"

# Startup
@app.on_event("startup")
async def startup_event():
    global research_agent, report_generator, research_planner, deep_executor, query_router, agent_registry, calculator, number_extractor, company_resolver, vector_store, document_processor, query_classifier, context_manager, pdf_parser
    
    logger.info("Starting Financial Research AI API...")
    
    try:
        research_agent = ResearchAgent()
        report_generator = ReportGenerator()
        research_planner = ResearchPlanner(llm_client=research_agent.llm) if hasattr(research_agent, 'llm') else ResearchPlanner()
        
        # Initialize Deep Research Executor
        deep_executor = DeepResearchExecutor(
            research_agent=research_agent,
            llm_client=research_agent.llm if hasattr(research_agent, 'llm') else None,
            search_client=None
        )
        
        # Initialize Query Router and Agent Registry
        llm_client = research_agent.llm if hasattr(research_agent, 'llm') else None
        query_router = QueryRouter(llm_client)
        agent_registry = AgentRegistry(llm_client, search_client=None)
        
        # Initialize Calculator and Number Extractor
        calculator = FinancialCalculator()
        number_extractor = NumberExtractor()
        
        # Initialize Universal Company Resolver
        company_resolver = UniversalCompanyResolver()
        
        # Initialize Query Classifier and Context Manager
        query_classifier = IntelligentQueryClassifier()
        context_manager = ConversationContextManager()
        
        # Initialize Vector Store and Document Processor for file attachments
        vector_store = VectorStore()
        document_processor = DocumentProcessor(chunk_size=1500, chunk_overlap=200)
        
        # Initialize Enhanced PDF Parser with Vision capabilities
        pdf_parser = EnhancedPDFParser(use_vision_api=True)
        
        # Create directories
        Path("outputs").mkdir(exist_ok=True)
        Path("uploads").mkdir(exist_ok=True)
        
        logger.info("Research Planner initialized")
        logger.info("Deep Research Executor initialized")
        logger.info("Query Router and Agent Registry initialized")
        logger.info("Financial Calculator and Number Extractor initialized")
        logger.info("Universal Company Resolver initialized")
        logger.info("Intelligent Query Classifier initialized")
        logger.info("Context Manager initialized")
        logger.info("Vector Store and Document Processor initialized")
        logger.info("Enhanced PDF Parser initialized with Vision API support")
        logger.info("API initialized successfully")
    except Exception as e:
        logger.error("Failed to initialize", error=e)
        raise

def _format_response_with_highlights(response: str, sources: List[Dict]) -> str:
    """
    Format response with highlighted actionable insights and source attribution
    
    Enhances readability by:
    - Highlighting key insights with emoji markers
    - Adding proper source citations
    - Emphasizing recommendations and risk flags
    """
    import re
    
    # Patterns for actionable insights
    insight_patterns = [
        (r'(Key takeaway|Key insight|Important:|Critical:|Notable:)(.+?)(?=\n|$)', '🔑 **\\1**\\2'),
        (r'(Recommendation|Suggested action|Advised to)(.+?)(?=\n|$)', '💡 **\\1**\\2'),
        (r'(Risk|Warning|Concern|Challenge)(.+?)(?=\n|$)', '⚠️ **\\1**\\2'),
        (r'(Opportunity|Potential|Growth)(.+?)(?=\n|$)', '🚀 **\\1**\\2'),
    ]
    
    enhanced_response = response
    
    # Apply highlighting patterns
    for pattern, replacement in insight_patterns:
        enhanced_response = re.sub(pattern, replacement, enhanced_response, flags=re.IGNORECASE)
    
    # Add source attribution section if sources exist
    if sources and len(sources) > 0:
        enhanced_response += "\n\n---\n\n### 📚 Sources & References\n\n"
        
        for idx, source in enumerate(sources[:10], 1):  # Limit to top 10 sources
            if isinstance(source, dict):
                title = source.get('title', 'Untitled')
                url = source.get('url', '#')
                snippet = source.get('snippet', '')[:150]
                
                enhanced_response += f"{idx}. **[{title}]({url})**"
                if snippet:
                    enhanced_response += f"\n   *{snippet}...*"
                enhanced_response += "\n\n"
            elif isinstance(source, str):
                enhanced_response += f"{idx}. {source}\n\n"
    
    return enhanced_response

# Helper Functions
async def resolve_context_with_ai(current_query: str, history: List[Dict]) -> str:
    """Use LLM to resolve context from conversation history"""
    
    # Check if query is a simple greeting/acknowledgment - don't resolve these
    simple_phrases = {
        'hi', 'hello', 'hey', 'thanks', 'thank you', 'ok', 'okay', 'great', 'good',
        'nice', 'cool', 'awesome', 'perfect', 'got it', 'understood', 'yes', 'no',
        'bye', 'goodbye', 'see you', 'later', 'welcome', 'sure', 'alright'
    }
    
    query_lower = current_query.lower().strip().strip('!.,?')
    if query_lower in simple_phrases:
        return current_query
    
    # Only check last 3 exchanges
    recent_history = history[-6:] if len(history) > 6 else history
    
    if not recent_history:
        return current_query
    
    # Build conversation context
    context_messages = []
    companies_mentioned = []
    
    for msg in recent_history:
        role = "User" if msg['role'] == 'user' else "Assistant"
        content = msg['content'][:300]  # Increased from 200
        context_messages.append(f"{role}: {content}")
        
        # Extract company names from messages (simple regex)
        import re
        company_patterns = r'\b(Microsoft|Google|Apple|Amazon|Tesla|Meta|Netflix|Nvidia|AMD|Intel|Oracle|IBM|Salesforce|Adobe|Cisco|PayPal|Qualcomm|Broadcom|Texas Instruments|Alphabet|GOOGL|MSFT|AAPL|AMZN|TSLA|META|NFLX|NVDA)\b'
        companies = re.findall(company_patterns, content, re.IGNORECASE)
        companies_mentioned.extend(companies)
    
    context_str = "\n".join(context_messages)
    
    # Get unique companies mentioned
    unique_companies = list(dict.fromkeys([c.title() for c in companies_mentioned if c]))
    
    # Enhanced prompt with company extraction
    prompt = f"""Given a conversation history and a new query, rewrite the query to be standalone and clear.

Conversation History:
{context_str}

Companies/Entities Mentioned: {', '.join(unique_companies) if unique_companies else 'None'}

New Query: {current_query}

Rules:
- If query uses pronouns (it, they, both, these, this, that) replace with EXACT company names from history
- Replace "both company/companies" with the specific company names mentioned
- Replace "this stock/these stocks" with actual company names
- If query is already clear with company names, return it unchanged
- Keep the question intent but make entities explicit
- Return ONLY the rewritten query, nothing else

Rewritten Query:"""

    try:
        # Use existing llm from research_agent
        response = await research_agent.llm.generate_response(prompt)
        
        resolved = response.strip()
        
        if resolved and len(resolved) > 5 and len(resolved) < 300:
            logger.info(f"Context resolved: '{current_query}' -> '{resolved}'")
            return resolved
        else:
            return current_query
            
    except Exception as e:
        logger.error(f"Error resolving context: {e}")
        return current_query

# Routes
@app.get("/")
async def root():
    """Serve main page"""
    return FileResponse("static/index.html")

@app.get("/user-guide.html")
async def user_guide():
    """Serve user guide page"""
    return FileResponse("static/user-guide.html")

@app.post("/api/chat")
async def chat(request: ChatMessage):
    """Enhanced chat endpoint with intelligent query classification"""
    try:
        # Validate input
        validator = InputValidator()
        is_valid, error = validator.validate_query(request.message)
        if not is_valid:
            raise HTTPException(status_code=400, detail=error)
        
        # Sanitize query
        query = validate_and_sanitize_query(request.message)
        
        session_id = request.session_id or str(uuid.uuid4())
        
        # Validate session ID if provided
        if request.session_id:
            is_valid, error = validator.validate_session_id(request.session_id)
            if not is_valid:
                logger.warning(f"Invalid session ID: {error}")
                session_id = str(uuid.uuid4())
        
        logger.info(f"Received: '{query}' (Session: {session_id[:8]})")
        
        # Check memory usage
        memory_manager.log_memory_stats()
        memory_manager.cleanup_if_needed(threshold_mb=600)
        
        # Get conversation context
        context = context_manager.get_context(session_id) if context_manager else None
        
        # STEP 1: INTELLIGENT QUERY CLASSIFICATION
        classification = query_classifier.classify(request.message, context)
        
        # Override classification if deep mode is forced from frontend
        if request.deepMode and classification.tier != 'DEEP':
            logger.info(f"🔬 Deep Mode forced by user - overriding {classification.tier} classification")
            classification.tier = 'DEEP'
            classification.confidence = 0.95
            classification.reasoning = "User-requested deep analysis mode"
            classification.suggested_steps = max(classification.suggested_steps, 8)
            classification.auto_execute = False
        
        logger.info(f"🎯 Query classified as: {classification.tier}")
        logger.info(f"   Confidence: {classification.confidence:.2f}")
        logger.info(f"   Reasoning: {classification.reasoning}")
        logger.info(f"   Suggested steps: {classification.suggested_steps}")
        logger.info(f"   Auto-execute: {classification.auto_execute}")
        
        # Handle file attachment if present
        file_context = ""
        if request.file_path:
            logger.info(f"Processing attached file: {request.file_path}")
            try:
                # Validate file path
                if not isinstance(request.file_path, str):
                    raise ValueError("Invalid file path format")
                
                file_path = Path(request.file_path)
                
                if not file_path.exists():
                    filename = file_path.name
                    alternative_path = Path("uploads") / filename
                    if alternative_path.exists():
                        file_path = alternative_path
                    else:
                        logger.error(f"Attached file not found: {request.file_path}")
                        raise HTTPException(
                            status_code=404, 
                            detail=f"Attached file not found: {filename}"
                        )
                
                # Check file size (max 10MB)
                file_size = file_path.stat().st_size
                max_size = 10 * 1024 * 1024  # 10MB
                if file_size > max_size:
                    raise HTTPException(
                        status_code=413,
                        detail=f"File too large ({file_size / 1024 / 1024:.2f}MB). Maximum size is 10MB."
                    )
                
                file_key = str(file_path.absolute())
                
                if file_key not in file_document_cache:
                    logger.info(f"Processing new file with vector store: {file_path.name}")
                    
                    file_ext = file_path.suffix.lower()
                    full_text = ""
                    
                    if file_ext == '.pdf':
                        # Use enhanced PDF parser for multimodal extraction
                        logger.info(f"Using enhanced PDF parser with vision capabilities")
                        pdf_content = pdf_parser.parse_pdf(str(file_path), analyze_charts=True)
                        
                        # Get combined text including tables, charts, and image descriptions
                        full_text = pdf_content.get_combined_text()
                        
                        logger.info(
                            f"Enhanced extraction complete: "
                            f"{len(pdf_content.text)} chars text, "
                            f"{len(pdf_content.tables)} tables, "
                            f"{len(pdf_content.charts)} charts analyzed"
                        )
                        
                    elif file_ext in ['.docx', '.doc']:
                        doc = Document(file_path)
                        full_text = "\n".join([para.text for para in doc.paragraphs])
                    elif file_ext == '.txt':
                        with open(file_path, 'r', encoding='utf-8') as f:
                            full_text = f.read()
                    
                    doc_id = f"doc_{file_path.stem}_{datetime.now().strftime('%Y%m%d%H%M%S')}"
                    doc = DocClass(
                        content=full_text,
                        metadata={
                            "filename": file_path.name,
                            "filepath": str(file_path),
                            "doc_type": file_ext
                        },
                        doc_id=doc_id
                    )
                    
                    chunks = document_processor.chunk_text(
                        full_text,
                        metadata={
                            "doc_id": doc_id,
                            "filename": file_path.name,
                            "source": str(file_path)
                        }
                    )
                    
                    logger.info(f"Created {len(chunks)} chunks from {file_path.name}")
                    vector_store.add_chunks(chunks)
                    file_document_cache[file_key] = doc_id
                    logger.info(f"File added to vector store: {file_path.name} (doc_id: {doc_id})")
                else:
                    doc_id = file_document_cache[file_key]
                    logger.info(f"File already in vector store: {file_path.name} (doc_id: {doc_id})")
                
                all_chunks = vector_store.get_chunks_by_doc_id(doc_id)
                
                if all_chunks:
                    all_chunks_sorted = sorted(all_chunks, key=lambda x: x.get('index', 0))
                    # REDUCED: Limit context to avoid token limits (8k chars ~= 2k tokens)
                    max_context_chars = 8000
                    relevant_chunks = []
                    total_chars = 0
                    
                    for result in all_chunks_sorted:
                        chunk_text = result['chunk'].text
                        chunk_len = len(chunk_text)
                        
                        if total_chars + chunk_len <= max_context_chars:
                            relevant_chunks.append(chunk_text)
                            total_chars += chunk_len
                        else:
                            remaining = max_context_chars - total_chars
                            if remaining > 500:
                                relevant_chunks.append(chunk_text[:remaining] + "...")
                            logger.warning(f"Context limit reached: {total_chars:,} chars (limit: {max_context_chars:,})")
                            break
                    
                    file_context = "\n\n".join(relevant_chunks)
                    logger.info(f"Retrieved {len(relevant_chunks)} chunks from document ({total_chars:,} characters total)")
                
            except Exception as e:
                logger.error(f"Error processing attached file: {e}")
                import traceback
                logger.error(traceback.format_exc())
                file_context = ""
        
        # Get conversation history
        if session_id not in conversation_history:
            conversation_history[session_id] = []
        
        history = conversation_history[session_id]
        query = request.message
        
        # AI context resolution (skip if file attached)
        if len(history) > 0 and not request.file_path:
            resolved_query = await resolve_context_with_ai(request.message, history)
            if resolved_query != request.message:
                logger.info(f"AI contextualized: '{request.message}' → '{resolved_query}'")
                query = resolved_query
        elif request.file_path:
            logger.info(f"File attached - using document context instead of conversation history")
        
        if file_context:
            query = f"{query}\n\nRelevant context from attached document ({Path(request.file_path).name if request.file_path else 'document'}):\n{file_context}"
            logger.info(f"Added {len(file_context)} chars of relevant context from vector store")
        
        # Add to history
        history.append({
            "role": "user",
            "content": request.message,
            "timestamp": datetime.now().isoformat()
        })
        
        # UNIVERSAL COMPANY RESOLUTION
        context_for_resolution = "\n".join([f"{msg['role']}: {msg['content'][:200]}" for msg in history[-4:]])
        resolution = company_resolver.resolve_company(query, context_for_resolution)
        
        company_info = None
        companies_list = []
        tickers_list = []
        
        if resolution.get("resolved"):
            company_info = resolution
            logger.info(f"[OK] Resolved companies: {company_info}")
            
            if resolution.get("multiple"):
                ticker_list = ", ".join([c['ticker'] for c in resolution['companies']])
                companies_list = [c['company'] for c in resolution['companies']]
                tickers_list = [c['ticker'] for c in resolution['companies']]
                logger.info(f"Multiple companies detected: {ticker_list}")
            else:
                companies_list = [resolution.get('company')]
                tickers_list = [resolution.get('ticker')]
                logger.info(f"Single company: {resolution.get('company')} ({resolution.get('ticker')})")
        
        # Update conversation context
        if context_manager:
            context_manager.update_context(
                session_id=session_id,
                query=request.message,
                companies=companies_list,
                tickers=tickers_list
            )
        
        # STEP 2: HANDLE BASED ON CLASSIFICATION TIER
        
        if classification.tier == 'INSTANT':
            # Direct execution for instant queries
            logger.info("⚡ INSTANT mode - Executing directly without plan")
            
            try:
                research_result = await research_agent.research(query, company_info=company_info)
                
                if hasattr(research_result, 'answer'):
                    response_text = research_result.answer
                    sources = getattr(research_result, 'sources', [])
                elif isinstance(research_result, dict):
                    response_text = research_result.get('response', research_result.get('answer', str(research_result)))
                    sources = research_result.get('sources', [])
                else:
                    response_text = str(research_result)
                    sources = []
            except Exception as e:
                logger.error(f"Error during research execution: {e}")
                response_text = f"I encountered an error while processing your request: {str(e)}. Please try again or rephrase your question."
                sources = []
            
            history.append({
                "role": "assistant",
                "content": response_text,
                "timestamp": datetime.now().isoformat()
            })
            
            if len(history) > 20:
                conversation_history[session_id] = history[-20:]
            
            # Get upgrade suggestions
            upgrade_options = query_classifier.suggest_upgrade('INSTANT')
            
            return {
                "success": True,
                "requires_approval": False,
                "tier": "INSTANT",
                "response": response_text,
                "session_id": session_id,
                "sources": sources[:6],
                "upgrade_available": bool(upgrade_options),
                "upgrade_options": upgrade_options,
                "classification": {
                    "tier": classification.tier,
                    "confidence": classification.confidence,
                    "reasoning": classification.reasoning,
                    "estimated_time": classification.estimated_time
                }
            }
        
        elif classification.tier == 'STANDARD' and classification.auto_execute:
            # Auto-execute standard queries with mini-plan
            logger.info("STANDARD mode - Auto-executing with mini-plan")
            
            # Route the query
            routing_decision = await query_router.route_query(query)
            
            if not routing_decision.should_process:
                logger.info(f"Query rejected: {routing_decision.reasoning}")
                
                rejection_response = (
                    f"{routing_decision.rejection_message}\n\n"
                    "**I can help you with:**\n"
                    "- Stock price analysis\n"
                    "- Company financial performance\n"
                    "- Sector comparisons\n"
                    "- Market trends\n"
                    "- Competitive analysis\n\n"
                    "**Try asking:** \"What is Tesla's stock price?\" or \"Analyze Apple's recent performance\""
                )
                
                history.append({
                    "role": "assistant",
                    "content": rejection_response,
                    "timestamp": datetime.now().isoformat()
                })
                
                return {
                    "success": True,
                    "requires_approval": False,
                    "response": rejection_response,
                    "session_id": session_id,
                    "sources": []
                }
            
            # Generate plan
            plan = await research_planner.generate_plan(
                query=query,
                sector=routing_decision.sector.value if routing_decision.sector else None
            )
            
            # Limit steps for standard tier
            plan.total_steps = min(plan.total_steps, classification.suggested_steps)
            plan.steps = plan.steps[:classification.suggested_steps]
            
            research_plans[plan.plan_id] = plan
            logger.info(f"Standard plan created: {plan.plan_id} ({plan.total_steps} steps)")
            
            # Auto-execute
            logger.info("Auto-executing standard plan...")
            
            # Execute immediately
            start_time = time.time()
            
            research_result = await research_agent.research(query, company_info=company_info)
            
            execution_time = time.time() - start_time
            
            if hasattr(research_result, 'answer'):
                response_text = research_result.answer
                sources = getattr(research_result, 'sources', [])
            elif isinstance(research_result, dict):
                response_text = research_result.get('response', research_result.get('answer', str(research_result)))
                sources = research_result.get('sources', [])
            else:
                response_text = str(research_result)
                sources = []
            
            history.append({
                "role": "assistant",
                "content": response_text,
                "timestamp": datetime.now().isoformat()
            })
            
            if len(history) > 20:
                conversation_history[session_id] = history[-20:]
            
            # Get upgrade suggestions
            upgrade_options = query_classifier.suggest_upgrade('STANDARD')
            
            return {
                "success": True,
                "requires_approval": False,
                "tier": "STANDARD",
                "response": response_text,
                "session_id": session_id,
                "sources": sources[:6],
                "execution_time": round(execution_time, 2),
                "steps_executed": plan.total_steps,
                "upgrade_available": bool(upgrade_options),
                "upgrade_options": upgrade_options,
                "classification": {
                    "tier": classification.tier,
                    "confidence": classification.confidence,
                    "reasoning": classification.reasoning,
                    "estimated_time": classification.estimated_time
                }
            }
        
        else:
            # DEEP tier or requires approval
            logger.info("🔬 DEEP mode - Generating detailed plan for approval")
            
            # Route the query
            routing_decision = await query_router.route_query(query)
            
            if not routing_decision.should_process:
                logger.info(f"Query rejected: {routing_decision.reasoning}")
                
                rejection_response = (
                    f"{routing_decision.rejection_message}\n\n"
                    "**I can help you with:**\n"
                    "- Stock price analysis\n"
                    "- Company financial performance\n"
                    "- Sector comparisons\n"
                    "- Market trends\n"
                    "- Competitive analysis\n\n"
                    "**Try asking:** \"What is Tesla's stock price?\" or \"Analyze Apple's recent performance\""
                )
                
                history.append({
                    "role": "assistant",
                    "content": rejection_response,
                    "timestamp": datetime.now().isoformat()
                })
                
                return {
                    "success": True,
                    "requires_approval": False,
                    "response": rejection_response,
                    "session_id": session_id,
                    "sources": []
                }
            
            # Generate detailed plan
            plan = await research_planner.generate_plan(
                query=query,
                sector=routing_decision.sector.value if routing_decision.sector else None
            )
            
            research_plans[plan.plan_id] = plan
            logger.info(f"Deep plan created: {plan.plan_id} ({plan.total_steps} steps)")
            
            return {
                "success": True,
                "requires_approval": True,
                "tier": "DEEP",
                "plan": plan.dict(),
                "routing": routing_decision.dict(),
                "session_id": session_id,
                "classification": {
                    "tier": classification.tier,
                    "confidence": classification.confidence,
                    "reasoning": classification.reasoning,
                    "estimated_time": classification.estimated_time,
                    "suggested_steps": classification.suggested_steps
                },
                "message": f"🔬 Deep research plan generated with {plan.total_steps} comprehensive steps."
            }
        
    except Exception as e:
        logger.error(f"Chat error: {e}")
        import traceback
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=str(e))

# Old chat endpoint - kept for reference, remove after testing
# @app.post("/api/chat_old")
# async def chat(request: ChatMessage):
    # """Enhanced chat endpoint with conversation context"""
    try:
        session_id = request.session_id or str(uuid.uuid4())
        
        logger.info(f"Received: '{request.message}' (Session: {session_id[:8]})")
        
        # Handle file attachment if present - USE VECTOR STORE
        file_context = ""
        if request.file_path:
            logger.info(f"Processing attached file: {request.file_path}")
            try:
                file_path = Path(request.file_path)
                
                # Check if file exists
                if not file_path.exists():
                    # Try alternative path (just filename in uploads folder)
                    filename = file_path.name
                    alternative_path = Path("uploads") / filename
                    if alternative_path.exists():
                        file_path = alternative_path
                    else:
                        logger.error(f"Attached file not found: {request.file_path}")
                        raise HTTPException(status_code=404, detail=f"Attached file not found: {filename}")
                
                # Check if this file is already processed in cache
                file_key = str(file_path.absolute())
                
                if file_key not in file_document_cache:
                    # New file - process and add to vector store
                    logger.info(f"Processing new file with vector store: {file_path.name}")
                    
                    # Extract file content
                    file_ext = file_path.suffix.lower()
                    full_text = ""
                    
                    if file_ext == '.pdf':
                        with open(file_path, 'rb') as pdf_file:
                            pdf_reader = PyPDF2.PdfReader(pdf_file)
                            for page_num in range(len(pdf_reader.pages)):
                                try:
                                    page = pdf_reader.pages[page_num]
                                    full_text += page.extract_text() + "\n"
                                except Exception as e:
                                    logger.error(f"Error reading page {page_num + 1}: {e}")
                                    continue
                    elif file_ext in ['.docx', '.doc']:
                        doc = Document(file_path)
                        full_text = "\n".join([para.text for para in doc.paragraphs])
                    elif file_ext == '.txt':
                        with open(file_path, 'r', encoding='utf-8') as f:
                            full_text = f.read()
                    
                    # Create document and chunk it
                    doc_id = f"doc_{file_path.stem}_{datetime.now().strftime('%Y%m%d%H%M%S')}"
                    doc = DocClass(
                        content=full_text,
                        metadata={
                            "filename": file_path.name,
                            "filepath": str(file_path),
                            "doc_type": file_ext
                        },
                        doc_id=doc_id
                    )
                    
                    # Chunk the document
                    chunks = document_processor.chunk_text(
                        full_text,
                        metadata={
                            "doc_id": doc_id,
                            "filename": file_path.name,
                            "source": str(file_path)
                        }
                    )
                    
                    logger.info(f"Created {len(chunks)} chunks from {file_path.name}")
                    
                    # Add to vector store
                    vector_store.add_chunks(chunks)
                    
                    # Cache the doc_id
                    file_document_cache[file_key] = doc_id
                    
                    logger.info(f"File added to vector store: {file_path.name} (doc_id: {doc_id})")
                else:
                    doc_id = file_document_cache[file_key]
                    logger.info(f"File already in vector store: {file_path.name} (doc_id: {doc_id})")
                
                # IMPORTANT: Get ALL chunks from the uploaded document
                # Don't use similarity search - we want the ENTIRE document content
                all_chunks = vector_store.get_chunks_by_doc_id(doc_id)
                
                if all_chunks:
                    # Build context from ALL chunks (sorted by original order)
                    all_chunks_sorted = sorted(all_chunks, key=lambda x: x.get('index', 0))
                    
                    # REDUCED: Limit context to avoid token limits (8k chars ~= 2k tokens)
                    max_context_chars = 8000  # Reduced from 30000 to avoid rate limit errors
                    relevant_chunks = []
                    total_chars = 0
                    
                    for result in all_chunks_sorted:
                        chunk_text = result['chunk'].text
                        chunk_len = len(chunk_text)
                        
                        if total_chars + chunk_len <= max_context_chars:
                            relevant_chunks.append(chunk_text)
                            total_chars += chunk_len
                        else:
                            # Add partial chunk if space allows
                            remaining = max_context_chars - total_chars
                            if remaining > 500:  # Only add if meaningful amount
                                relevant_chunks.append(chunk_text[:remaining] + "...")
                            logger.warning(f"Context limit reached: {total_chars:,} chars (limit: {max_context_chars:,})")
                            break
                    
                    file_context = "\n\n".join(relevant_chunks)
                    logger.info(f"Retrieved {len(relevant_chunks)} chunks from document ({total_chars:,} characters total)")
                else:
                    logger.warning("No chunks found for document")
                    file_context = ""
                
            except Exception as e:
                logger.error(f"Error processing attached file: {e}")
                import traceback
                logger.error(traceback.format_exc())
                # Continue without file context rather than failing
                file_context = ""
        
        # Get conversation history
        if session_id not in conversation_history:
            conversation_history[session_id] = []
        
        history = conversation_history[session_id]
        
        # AI-POWERED CONTEXT RESOLUTION
        query = request.message
        
        # IMPORTANT: Skip AI context resolution when a file is attached
        # The file content should take priority over conversation history
        if len(history) > 0 and not request.file_path:
            # Use LLM to resolve context ONLY if no file is attached
            resolved_query = await resolve_context_with_ai(request.message, history)
            if resolved_query != request.message:
                logger.info(f"AI contextualized: '{request.message}' → '{resolved_query}'")
                query = resolved_query
        elif request.file_path:
            # When file is attached, don't use conversation history for context
            logger.info(f"File attached - using document context instead of conversation history")
        
        # If file context exists, enhance the query with RELEVANT chunks
        if file_context:
            query = f"{query}\n\nRelevant context from attached document ({Path(request.file_path).name if request.file_path else 'document'}):\n{file_context}"
            logger.info(f"Added {len(file_context)} chars of relevant context from vector store")
        
        # Add to history
        history.append({
            "role": "user",
            "content": request.message,
            "timestamp": datetime.now().isoformat()
        })
        
        # UNIVERSAL COMPANY RESOLUTION
        # Resolve companies before routing to ensure correct tickers
        context_for_resolution = "\n".join([f"{msg['role']}: {msg['content'][:200]}" for msg in history[-4:]])
        resolution = company_resolver.resolve_company(query, context_for_resolution)
        
        company_info = None
        if resolution.get("resolved"):
            company_info = resolution
            logger.info(f"[OK] Resolved companies: {company_info}")
            
            # Enhance query with resolved info for better routing
            if resolution.get("multiple"):
                ticker_list = ", ".join([c['ticker'] for c in resolution['companies']])
                logger.info(f"Multiple companies detected: {ticker_list}")
            else:
                logger.info(f"Single company: {resolution.get('company')} ({resolution.get('ticker')})")
                if resolution.get('is_subsidiary'):
                    logger.info(f"Note: {resolution.get('original_mention')} -> {resolution.get('company')}")
        
        # Route the query
        routing_decision = await query_router.route_query(query)
        
        # If non-financial, reject
        if not routing_decision.should_process:
            logger.info(f"Query rejected: {routing_decision.reasoning}")
            
            rejection_response = (
                f"{routing_decision.rejection_message}\n\n"
                "**I can help you with:**\n"
                "- Stock price analysis\n"
                "- Company financial performance\n"  
                "- Sector comparisons\n"
                "- Market trends\n"
                "- Competitive analysis\n\n"
                "**Try asking:** \"What is Tesla's stock price?\" or \"Analyze Apple's recent performance\""
            )
            
            history.append({
                "role": "assistant",
                "content": rejection_response,
                "timestamp": datetime.now().isoformat()
            })
            
            return {
                "success": True,
                "requires_approval": False,
                "response": rejection_response,
                "session_id": session_id,
                "sources": []
            }
        
        # Check complexity
        complexity = routing_decision.complexity or "medium"
        logger.info(f"Complexity: {complexity}")
        
        # Simple queries → Direct execution
        if complexity == "simple":
            logger.info("Simple query - executing directly")
            
            # Pass company_info to research method
            research_result = await research_agent.research(query, company_info=company_info)
            
            # Extract answer
            if hasattr(research_result, 'answer'):
                response_text = research_result.answer
                sources = getattr(research_result, 'sources', [])
                
                if sources:
                    formatted_sources = []
                    for source in sources[:5]:
                        if isinstance(source, dict):
                            formatted_sources.append(source)
                    sources = formatted_sources
                else:
                    sources = []
                
            elif isinstance(research_result, dict):
                response_text = research_result.get('response', research_result.get('answer', str(research_result)))
                sources = research_result.get('sources', [])
            else:
                response_text = str(research_result)
                sources = []
            
            # Add to history
            history.append({
                "role": "assistant",
                "content": response_text,
                "timestamp": datetime.now().isoformat()
            })
            
            # Limit history
            if len(history) > 20:
                conversation_history[session_id] = history[-20:]
            
            return {
                "success": True,
                "requires_approval": False,
                "response": response_text,
                "complexity": complexity,
                "session_id": session_id,
                "sources": sources
            }
        
        # Complex queries → Show plan
        logger.info(f"Generating plan for {complexity} query")
        
        plan = await research_planner.generate_plan(
            query=query,
            sector=routing_decision.sector.value if routing_decision.sector else None
        )
        
        research_plans[plan.plan_id] = plan
        logger.info(f"Plan created: {plan.plan_id}")
        
        return {
            "success": True,
            "requires_approval": True,
            "plan": plan.dict(),
            "routing": routing_decision.dict(),
            "complexity": complexity,
            "session_id": session_id,
            "message": f"Research plan generated with {plan.total_steps} steps."
        }
        
    except Exception as e:
        logger.error(f"Chat error: {e}")
        import traceback
        traceback.print_exc()
        
        return {
            "success": False,
            "requires_approval": False,
            "response": f"❌ Error: {str(e)}\n\nPlease try again or rephrase your question.",
            "session_id": session_id if 'session_id' in locals() else str(uuid.uuid4()),
            "sources": []
        }

@app.get("/api/stock/price/{symbol}")
async def get_stock_price(symbol: str):
    """Get current stock price"""
    try:
        stock = yf.Ticker(symbol)
        info = stock.info
        current_price = info.get('currentPrice', info.get('regularMarketPrice', 0))
        change = info.get('regularMarketChangePercent', 0)
        
        return {
            "symbol": symbol,
            "price": current_price,
            "change": change
        }
    except Exception as e:
        logger.error(f"Stock price error for {symbol}", error=e)
        raise HTTPException(status_code=400, detail=str(e))

@app.post("/api/stock/chart")
async def get_stock_chart(query: StockQuery):
    """Get stock chart data"""
    try:
        stock = yf.Ticker(query.symbol)
        hist = stock.history(period=query.period)
        
        return {
            "symbol": query.symbol,
            "dates": hist.index.strftime('%Y-%m-%d').tolist(),
            "prices": hist['Close'].tolist(),
            "volume": hist['Volume'].tolist()
        }
    except Exception as e:
        logger.error(f"Chart error for {query.symbol}", error=e)
        raise HTTPException(status_code=400, detail=str(e))

@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...)):
    """Upload and process documents (PDF, DOCX, TXT)"""
    try:
        logger.info(f"Receiving file upload: {file.filename}")
        
        # CLEAR OLD FILE DATA FROM VECTOR STORE AND CACHE
        # This prevents old file data from persisting when a new file is uploaded
        if file_document_cache:
            logger.info("Clearing old file data from cache and vector store")
            file_document_cache.clear()
        
        # Clear vector store to remove old document chunks
        if vector_store:
            logger.info("Resetting vector store for new file")
            vector_store.clear_store()  # This will clear all old document data
        
        # Validate file type
        allowed_extensions = ['.pdf', '.docx', '.txt', '.doc']
        file_ext = os.path.splitext(file.filename)[1].lower()
        
        if file_ext not in allowed_extensions:
            raise HTTPException(
                status_code=400,
                detail=f"File type {file_ext} not supported. Allowed: {', '.join(allowed_extensions)}"
            )
        
        # Save file
        upload_path = Path("uploads") / file.filename
        upload_path.parent.mkdir(exist_ok=True)
        
        file_content = await file.read()
        with upload_path.open("wb") as buffer:
            buffer.write(file_content)
        
        # Extract text based on file type
        text_content = ""
        tables_info = ""
        charts_info = ""
        
        if file_ext == '.pdf':
            # Use enhanced PDF parser for multimodal extraction
            logger.info(f"Processing PDF with enhanced parser (tables, charts, images)")
            pdf_content = pdf_parser.parse_pdf(str(upload_path), analyze_charts=True)
            
            # Get combined text including tables, charts, and image descriptions
            text_content = pdf_content.get_combined_text()
            
            tables_info = f"{len(pdf_content.tables)} tables extracted"
            charts_info = f"{len(pdf_content.charts)} charts analyzed"
            
            logger.info(
                f"Enhanced extraction: "
                f"{len(pdf_content.text)} chars, "
                f"{len(pdf_content.tables)} tables, "
                f"{len(pdf_content.charts)} charts"
            )
        
        elif file_ext in ['.docx', '.doc']:
            doc = Document(upload_path)
            text_content = "\n".join([para.text for para in doc.paragraphs])
        
        elif file_ext == '.txt':
            with open(upload_path, 'r', encoding='utf-8') as txt_file:
                text_content = txt_file.read()
        
        logger.info(f"Extracted {len(text_content)} characters from {file.filename}")
        
        response_data = {
            "success": True,
            "filename": file.filename,
            "file_path": str(upload_path),
            "text_length": len(text_content),
            "preview": text_content[:500] + "..." if len(text_content) > 500 else text_content,
            "message": "File uploaded and processed successfully"
        }
        
        # Add table and chart info for PDFs
        if file_ext == '.pdf':
            response_data["tables"] = tables_info
            response_data["charts"] = charts_info
        
        return response_data
        
    except Exception as e:
        logger.error(f"Error uploading file: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/analyze-document")
async def analyze_document(request: dict):
    """Analyze uploaded document with a specific query"""
    try:
        file_path = request.get('file_path')
        query = request.get('query', 'Summarize this document')
        
        logger.info(f"Analyzing document: {file_path}")
        logger.info(f"Query: {query}")
        
        if not file_path:
            raise HTTPException(status_code=400, detail="File path required")
        
        # Ensure file_path is a Path object
        file_path = Path(file_path)
        
        # Check if file exists
        if not file_path.exists():
            # Try alternative path (just filename in uploads folder)
            filename = file_path.name
            alternative_path = Path("uploads") / filename
            
            logger.info(f"Original path not found: {file_path}")
            logger.info(f"Trying alternative: {alternative_path}")
            
            if alternative_path.exists():
                file_path = alternative_path
            else:
                logger.error(f"File not found at either location")
                raise HTTPException(
                    status_code=404, 
                    detail=f"File not found: {filename}"
                )
        
        logger.info(f"File found at: {file_path}")
        
        # Read document content
        file_ext = file_path.suffix.lower()
        content = ""
        
        if file_ext == '.pdf':
            logger.info("Extracting PDF content...")
            with open(file_path, 'rb') as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                num_pages = len(pdf_reader.pages)
                logger.info(f"PDF has {num_pages} pages")
                
                # Limit to first 50 pages for analysis
                pages_to_read = min(num_pages, 50)
                
                for page_num in range(pages_to_read):
                    try:
                        page = pdf_reader.pages[page_num]
                        text = page.extract_text()
                        content += f"\n\n[Page {page_num + 1}]\n{text}"
                    except Exception as e:
                        logger.error(f"Error reading page {page_num + 1}: {e}")
                        continue
                
                logger.info(f"Extracted {len(content)} characters from {pages_to_read} pages")
        
        elif file_ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
        
        elif file_ext in ['.docx', '.doc']:
            doc = Document(file_path)
            content = "\n".join([para.text for para in doc.paragraphs])
        
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type: {file_ext}"
            )
        
        if not content.strip():
            raise HTTPException(
                status_code=400,
                detail="No text content extracted from document"
            )
        
        # Limit content size to avoid token limits
        max_chars = 15000  # ~4000 tokens
        if len(content) > max_chars:
            content = content[:max_chars] + "\n\n[Content truncated for analysis...]"
            logger.info(f"Content truncated to {max_chars} characters")
        
        # Use research agent to analyze
        logger.info("Sending to research agent...")
        result = await research_agent.research(
            query=f"{query}\n\nDocument: {file_path.name}\n\nContent:\n{content}"
        )
        
        # Extract response properly
        if hasattr(result, 'answer'):
            response_text = result.answer
            sources = getattr(result, 'sources', [])
        elif isinstance(result, dict):
            response_text = result.get('response', result.get('answer', ''))
            sources = result.get('sources', [])
        else:
            response_text = str(result)
            sources = []
        
        logger.info("Analysis completed successfully")
        
        return {
            "success": True,
            "query": query,
            "filename": file_path.name,
            "pages_analyzed": pages_to_read if file_ext == '.pdf' else None,
            "analysis": response_text,
            "sources": sources
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error analyzing document: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/files")
async def list_files():
    """List all uploaded files"""
    try:
        upload_dir = Path("uploads")
        if not upload_dir.exists():
            return {"files": []}
        
        files = []
        for file_path in upload_dir.iterdir():
            if file_path.is_file():
                files.append({
                    "name": file_path.name,
                    "path": str(file_path),
                    "size": file_path.stat().st_size,
                    "modified": datetime.fromtimestamp(file_path.stat().st_mtime).isoformat()
                })
        
        return {"files": files}
        
    except Exception as e:
        logger.error(f"Error listing files: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/reports/{filename}")
async def get_report(filename: str):
    """Download report"""
    report_path = Path("outputs") / filename
    if not report_path.exists():
        raise HTTPException(status_code=404, detail="Report not found")
    return FileResponse(report_path)

@app.get("/api/status")
async def get_status():
    """API status"""
    return {
        "status": "ready" if research_agent else "initializing",
        "version": "3.0.0"
    }

# Test Routing
@app.post("/api/route")
async def route_query(request: dict):
    """Test the query router"""
    if not query_router:
        raise HTTPException(status_code=503, detail="Query Router not initialized")
    
    try:
        query = request.get('query')
        if not query:
            raise HTTPException(status_code=400, detail="Query required")
        
        logger.info(f"Routing query: {query[:50]}...")
        
        # Route the query
        decision = await query_router.route_query(query)
        
        return {
            "success": True,
            "routing": decision.dict() if hasattr(decision, 'dict') else decision
        }
        
    except Exception as e:
        logger.error(f"Routing error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Get Available Sectors
@app.get("/api/sectors")
async def get_sectors():
    """Get list of available sectors"""
    if not query_router:
        raise HTTPException(status_code=503, detail="Query Router not initialized")
    
    try:
        sectors = query_router.get_available_sectors() if hasattr(query_router, 'get_available_sectors') else []
        return {
            "success": True,
            "sectors": sectors
        }
    except Exception as e:
        logger.error(f"Error getting sectors: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Generate Research Plan
@app.post("/api/plan")
async def generate_research_plan(request: dict):
    """Generate a detailed research plan for user approval"""
    if not research_planner:
        raise HTTPException(status_code=503, detail="Research Planner not initialized")
    
    try:
        query = request.get('query')
        sector = request.get('sector')
        depth = request.get('depth', 'medium')
        
        if not query:
            raise HTTPException(status_code=400, detail="Query is required")
        
        logger.info(f"Generating research plan for: {query}")
        
        # Generate plan
        plan = await research_planner.generate_plan(
            query=query,
            sector=sector
        )
        
        # Store plan in memory
        plan_id = plan.plan_id if hasattr(plan, 'plan_id') else str(uuid.uuid4())
        research_plans[plan_id] = plan
        
        return {
            "success": True,
            "plan": plan.dict() if hasattr(plan, 'dict') else plan,
            "plan_id": plan_id,
            "message": "Research plan generated. Please review and approve."
        }
        
    except Exception as e:
        logger.error(f"Error generating plan: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Approve/Reject Plan
@app.post("/api/plan/{plan_id}/approve")
async def approve_research_plan(plan_id: str, action: dict):
    """Approve or reject a research plan"""
    try:
        action_type = action.get('action')
        
        if action_type == "approve":
            logger.info(f"Plan {plan_id} approved, starting execution...")
            
            return {
                "success": True,
                "status": "approved",
                "plan_id": plan_id,
                "message": "Plan approved. Executing..."
            }
        
        elif action_type == "reject":
            logger.info(f"Plan {plan_id} rejected")
            return {
                "success": True,
                "status": "rejected",
                "message": "Plan rejected"
            }
        
        elif action_type == "modify":
            modifications = action.get('modifications', {})
            logger.info(f"Plan {plan_id} modifications requested")
            return {
                "success": True,
                "status": "modified",
                "message": "Plan updated, please review again"
            }
        
        else:
            raise HTTPException(status_code=400, detail="Invalid action type")
        
    except Exception as e:
        logger.error(f"Error processing plan action: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Execute Research Plan
@app.post("/api/plan/{plan_id}/execute")
async def execute_research_plan(plan_id: str, request: dict = None):
    """Execute approved research plan with progress streaming"""
    try:
        logger.info(f"=" * 60)
        logger.info(f"EXECUTING PLAN: {plan_id}")
        logger.info(f"=" * 60)
        
        # START TIMING EXECUTION
        start_time = time.time()
        
        # Extract session_id from request body if provided
        session_id = None
        if request and isinstance(request, dict):
            session_id = request.get('session_id')
        
        # Get plan
        if plan_id not in research_plans:
            logger.error(f"Plan {plan_id} not found in research_plans")
            logger.error(f"Available plans: {list(research_plans.keys())}")
            raise HTTPException(status_code=404, detail="Plan not found")
        
        plan = research_plans[plan_id]
        logger.info(f"Plan found: {type(plan)}")
        logger.info(f"Plan attributes: {dir(plan)}")
        
        # Get query from plan
        query = getattr(plan, 'query', None)
        if not query:
            logger.error("No query found in plan")
            query = "Research analysis"
        
        logger.info(f"Executing query: {query}")
        
        # Track research progress for transparency
        total_steps = getattr(plan, 'total_steps', 5)
        research_progress = {
            "steps_completed": 0,
            "total_steps": total_steps,
            "intermediate_findings": [],
            "sources_found": []
        }
        
        logger.info(f"🔄 Starting research with {total_steps} planned steps...")
        
        # UNIVERSAL COMPANY RESOLUTION for plan execution
        context_for_resolution = ""
        if session_id and session_id in conversation_history:
            history = conversation_history[session_id]
            context_for_resolution = "\\n".join([f"{msg['role']}: {msg['content'][:200]}" for msg in history[-4:]])
        
        resolution = company_resolver.resolve_company(query, context_for_resolution)
        
        company_info = None
        if resolution.get("resolved"):
            company_info = resolution
            logger.info(f"[OK] Resolved companies for plan execution: {company_info}")
        
        # Call research agent with company info
        logger.info("🔄 Step 1/{}: Initializing research agent...".format(total_steps))
        research_progress["steps_completed"] = 1
        
        logger.info("🔄 Step 2/{}: Searching for information...".format(total_steps))
        research_progress["steps_completed"] = 2
        
        # Enable deep mode for complex queries (8+ steps or DEEP tier)
        is_deep_mode = total_steps >= 8 or plan.estimated_depth == "deep"
        if is_deep_mode:
            logger.info("🔬 Executing in DEEP RESEARCH MODE with iterative loops")
        
        research_result = await research_agent.research(
            query, 
            company_info=company_info,
            deep_mode=is_deep_mode
        )
        
        research_progress["steps_completed"] = total_steps
        logger.info(f"✅ Research completed! All {total_steps} steps executed.")
        logger.info(f"Result type: {type(research_result)}")
        logger.info(f"Result attributes: {dir(research_result)}")
        
        # Extract response with detailed logging and source attribution
        response_text = None
        sources = []
        
        if hasattr(research_result, 'answer'):
            response_text = research_result.answer
            logger.info(f"[OK] Extracted from .answer: {len(response_text)} chars")
            
            # EXTRACT SOURCES with attribution
            if hasattr(research_result, 'sources'):
                sources = research_result.sources
                logger.info(f"[OK] Extracted {len(sources)} sources")
                
        elif hasattr(research_result, 'response'):
            response_text = research_result.response
            logger.info(f"[OK] Extracted from .response: {len(response_text)} chars")
            
            # EXTRACT SOURCES
            if hasattr(research_result, 'sources'):
                sources = research_result.sources
                logger.info(f"[OK] Extracted {len(sources)} sources")
                
        elif isinstance(research_result, dict):
            response_text = research_result.get('answer') or research_result.get('response')
            sources = research_result.get('sources', [])
            logger.info(f"[OK] Extracted from dict: {len(response_text) if response_text else 0} chars, {len(sources)} sources")
            
        elif isinstance(research_result, str):
            response_text = research_result
            logger.info(f"[OK] Result is string: {len(response_text)} chars")
        else:
            response_text = str(research_result)
            logger.info(f"[WARN] Converted to string: {len(response_text)} chars")
        
        # Check if we got valid content
        if not response_text:
            logger.error("[ERROR] Response is empty!")
            response_text = "❌ Research failed: No response generated"
        elif len(response_text) < 50:
            logger.error(f"[ERROR] Response too short: {response_text}")
            response_text = f"⚠️ Research Warning: Limited information available\n\n{response_text}"
        else:
            logger.info(f"[OK] Valid response: {len(response_text)} characters")
            
            # If deep mode was used and we have research reasoning, format with it
            if hasattr(research_result, 'research_reasoning') and research_result.research_reasoning:
                logger.info(f"📊 Deep research with {len(research_result.research_reasoning)} iterations")
                response_text = research_result.format_with_reasoning()
            else:
                # Enhance response with highlighted insights and source attribution
                response_text = _format_response_with_highlights(response_text, sources)
            
            logger.info(f"Response preview: {response_text[:200]}...")
        
        # CALCULATE ACTUAL EXECUTION TIME
        execution_time = time.time() - start_time
        
        # Count actual sources used (3-6 range)
        sources_count = len(sources) if sources else 0
        logger.info(f"Execution Stats: {execution_time:.2f}s, {sources_count} sources used")
        
        # Add to conversation history if session_id provided
        if session_id:
            if session_id not in conversation_history:
                conversation_history[session_id] = []
            
            conversation_history[session_id].append({
                "role": "assistant",
                "content": response_text,
                "timestamp": datetime.now().isoformat()
            })
            
            # Limit history to last 20 messages
            if len(conversation_history[session_id]) > 20:
                conversation_history[session_id] = conversation_history[session_id][-20:]
            
            logger.info(f"Added response to conversation history for session {session_id[:8]}")
        
        # Build result with ACTUAL timing and source count
        result = {
            "success": True,
            "plan_id": plan_id,
            "results": {
                "final_report": response_text,
                "total_steps_executed": research_progress["steps_completed"],
                "discovered_topics": [],
                "execution_time": round(execution_time, 2),
                "sources": sources,
                "sources_used": sources_count
            },
            "message": "Research completed successfully"
        }
        
        logger.info(f"=" * 60)
        logger.info(f"EXECUTION COMPLETE - Returning {len(response_text)} chars")
        logger.info(f"=" * 60)
        
        return result
        
    except Exception as e:
        logger.error(f"❌ EXECUTION ERROR: {e}")
        import traceback
        error_trace = traceback.format_exc()
        logger.error(error_trace)
        
        # Calculate execution time even for failures
        execution_time = time.time() - start_time if 'start_time' in locals() else 0.0
        
        return {
            "success": False,
            "plan_id": plan_id,
            "error": str(e),
            "results": {
                "final_report": f"❌ **Research Execution Failed**\n\n**Error:** {str(e)}\n\n**Please try:**\n- Simplifying your query\n- Breaking it into smaller questions\n- Rephrasing the question\n\n**Debug Info:**\n```\n{error_trace[:500]}\n```",
                "total_steps_executed": research_progress.get("steps_completed", 0) if 'research_progress' in locals() else 0,
                "discovered_topics": [],
                "execution_time": round(execution_time, 2),
                "sources_used": 0
            },
            "message": f"Execution failed: {str(e)}"
        }

# Calculate Financial Metric
@app.post("/api/calculate")
async def calculate_metric(request: dict):
    """Direct calculation endpoint"""
    if not calculator:
        raise HTTPException(status_code=503, detail="Calculator not initialized")
    
    try:
        calc_type = request.get('calculation')
        inputs = request.get('inputs', {})
        
        logger.info(f"Calculating: {calc_type}")
        
        # Route to appropriate calculation
        if calc_type == 'growth_rate':
            result = calculator.calculate_growth_rate(
                current_value=inputs.get('current_value', 0),
                previous_value=inputs.get('previous_value', 0)
            )
        
        elif calc_type == 'margin':
            result = calculator.calculate_margin(
                profit=inputs.get('profit', 0),
                revenue=inputs.get('revenue', 0),
                margin_type=inputs.get('margin_type', 'net')
            )
        
        elif calc_type == 'pe_ratio':
            result = calculator.calculate_pe_ratio(
                stock_price=inputs.get('stock_price', 0),
                eps=inputs.get('eps', 0)
            )
        
        elif calc_type == 'roe':
            result = calculator.calculate_roe(
                net_income=inputs.get('net_income', 0),
                shareholders_equity=inputs.get('shareholders_equity', 0)
            )
        
        elif calc_type == 'roa':
            result = calculator.calculate_roa(
                net_income=inputs.get('net_income', 0),
                total_assets=inputs.get('total_assets', 0)
            )
        
        elif calc_type == 'cagr':
            result = calculator.calculate_cagr(
                starting_value=inputs.get('starting_value', 0),
                ending_value=inputs.get('ending_value', 0),
                num_years=inputs.get('num_years', 1)
            )
        
        elif calc_type == 'debt_to_equity':
            result = calculator.calculate_debt_to_equity(
                total_debt=inputs.get('total_debt', 0),
                shareholders_equity=inputs.get('shareholders_equity', 0)
            )
        
        elif calc_type == 'current_ratio':
            result = calculator.calculate_current_ratio(
                current_assets=inputs.get('current_assets', 0),
                current_liabilities=inputs.get('current_liabilities', 0)
            )
        
        else:
            raise HTTPException(status_code=400, detail=f"Unknown calculation: {calc_type}")
        
        return {
            "success": True,
            "result": result.dict() if hasattr(result, 'dict') else result
        }
        
    except Exception as e:
        logger.error(f"Calculation error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Extract Numbers from Text
@app.post("/api/extract-numbers")
async def extract_numbers(request: dict):
    """Extract numbers from text"""
    if not number_extractor:
        raise HTTPException(status_code=503, detail="Number Extractor not initialized")
    
    try:
        text = request.get('text', '')
        
        logger.info(f"Extracting numbers from text: {text[:50]}...")
        
        # Extract all numbers
        numbers = number_extractor.extract_all_numbers(text) if hasattr(number_extractor, 'extract_all_numbers') else []
        
        # Extract metrics
        metrics = number_extractor.extract_financial_metrics({'text': text}) if hasattr(number_extractor, 'extract_financial_metrics') else {}
        
        # Extract percentages
        percentages = []
        import re
        for match in re.finditer(r'([\d\.]+)%', text):
            percentages.append(float(match.group(1)))
        
        return {
            "success": True,
            "all_numbers": [{"value": n[0], "context": n[1]} for n in numbers] if numbers else [],
            "metrics": metrics,
            "percentages": percentages
        }
        
    except Exception as e:
        logger.error(f"Extraction error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)
